import os
import shutil
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

# 获取当前文件夹路径
Current_Folder_path = os.getcwd()
# 获取路径下的所有文件
file_list = os.listdir(Current_Folder_path)
file_list = [i for i in file_list if i.endswith('g')]

# 创建成果输出文件夹
if not os.path.exists(Current_Folder_path + '\\NewFolder'):
    os.mkdir(os.path.join(Current_Folder_path, 'NewFolder'))
for jpg in file_list:
    documentfile = Document(jpg[10:15] + ".docx")
    jpg_path = os.path.join(Current_Folder_path, jpg)
    # 指定段落（第一段）插入图片
    documentfile.paragraphs[0].insert_paragraph_before('')
    pic1 = documentfile.paragraphs[0].add_run().add_picture(jpg_path, width=Inches(6))  # 添加图片，设置宽度（长度等比例放大缩小）
    paragraph = documentfile.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中设置
    # word文档保存
    documentfile.save(os.path.join(Current_Folder_path, 'NewFolder', jpg[10:15] + ".docx"))
    # print(Folder,'文档生成完毕')

