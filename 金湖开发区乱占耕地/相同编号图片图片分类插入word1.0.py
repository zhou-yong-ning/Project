import os
import shutil
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

# 获取当前文件夹路径
Current_Folder_path = os.getcwd()
# 获取路径下的所有文件
file_list = os.listdir(Current_Folder_path)
file_list = [i for i in file_list if i.endswith('g') or i.endswith('G')]
# 截取前17位字符串
filename0 = [i[0:5] for i in file_list]
# 去除列表中的重复项（顺序改变）
filename0 = list(set(filename0))
# 根据编号建立文件夹
for j in filename0:
    if not os.path.exists(os.path.join(Current_Folder_path, j)):
        os.mkdir(os.path.join(Current_Folder_path, j))
# 移动文件
for k in file_list:
    shutil.move(os.path.join(Current_Folder_path, k), os.path.join(Current_Folder_path, k[0:5], k))
# 图片插入word
# 获取当前文件夹路径
# 遍历当前文件夹下的文件
file_list = os.listdir(Current_Folder_path)
# 创建成果输出文件夹
if not os.path.exists(Current_Folder_path + '\\NewFolder'):
    os.mkdir(os.path.join(Current_Folder_path, 'NewFolder'))
Folder_List = [i for i in file_list if not "." in i]
for Folder in Folder_List:
    # 获取子文件夹内的文件
    file_list1 = os.listdir(os.path.join(Current_Folder_path, Folder))
    # 提取后缀为JPG格式的文件
    JPG_list = [i for i in file_list1 if i.endswith('g') or i.endswith('G') ]
    # 判断文件夹是否含有两张图片
    # 创建一个word对象
    document = Document()
    for jpg in JPG_list:
        # 插入图片必须绝对路径
        jpg_path = os.path.join(Current_Folder_path,Folder,jpg)
        # 插入图片1
        pic1 = document.add_picture(jpg_path, width=Inches(5))  # 添加图片，设置宽度（长度等比例放大缩小）
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中设置
        # document.paragraphs[0].paragraph_format.space_before = Pt(20)  # 段前
        document.paragraphs[-1].paragraph_format.space_after = Pt(100)  # 段后
        # word文档保存
    document.save(os.path.join(Current_Folder_path,'NewFolder',f'{Folder}.docx'))
    print(Folder,'文档生成完毕')
