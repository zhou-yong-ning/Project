import os
from multiprocessing import Process
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

if __name__ == '__main__':
    # 获取当前文件夹路径
    Current_Folder_path = os.getcwd()
    # 遍历当前文件夹下的文件
    file_list = os.listdir(Current_Folder_path)
    # 创建成果输出文件夹
    if not os.path.exists(Current_Folder_path + '\\NewFolder'):
        os.mkdir(os.path.join(Current_Folder_path, 'NewFolder'))
    Folder_List = [i for i in file_list if not "." in i]
    for Folder in Folder_List:
        # 获取子文件夹内的文件
        file_list1 = os.listdir(os.path.join(Current_Folder_path, Folder))
        # 提取后缀为JPG格式的文件
        JPG_list = [i for i in file_list1 if i.endswith('.jpg')]
        # 判断文件夹是否含有两张图片
        if len(JPG_list) < 2:
            continue
        else:
            jpg1 = os.path.join(Current_Folder_path,Folder,JPG_list[0])
            jpg2 = os.path.join(Current_Folder_path,Folder,JPG_list[1])
            # 创建一个word对象
            document = Document()
            # 插入图片1
            pic1 = document.add_picture(jpg1, width=Inches(4))  # 添加图片，设置宽度（长度等比例放大缩小）
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中设置
            # 插入图片2
            pic2 = document.add_picture(jpg2, width=Inches(4))  # 添加图片，设置宽度（长度等比例放大缩小）
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 图片居中设置
            # 设置段前段后距离
            document.paragraphs[0].paragraph_format.space_before = Pt(30)  # 段前
            document.paragraphs[0].paragraph_format.space_after = Pt(100)  # 段后
            # word文档保存
            document.save(os.path.join(Current_Folder_path,'NewFolder',f'{Folder}.docx'))
            print(Folder,'身份证word生成完毕')
